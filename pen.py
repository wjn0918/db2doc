"""
将数据输出到word表格中
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm,Pt,RGBColor
from docx.shared import Pt
from docx.oxml.ns import qn





class MyPen():

    def __init__(self, file_path, file_type="word"):
        """
        初始化myPen
        :param file_type:输出文件类型
        """
        self.file_type = file_type
        self.file_path = file_path

        pass


    def draw(self,data):
        if self.file_type == 'word':
            self.__pen2word__(data)

    def __pen2word__(self,data):
        doc = Document()
        # 添加内容
        for index, table_name in enumerate(data['table_name']):
            table_comment = data['table_comment'][index]
            table_name = table_comment + "(" + table_name + ")"
            self.__addParagraph__(doc, table_name)
            column_index = data['column_index'][index]
            column_name = data['column_name'][index]
            column_type = data['column_type'][index]
            column_comment = data['column_comment'][index]
            talbe_content = zip(column_index, column_comment, column_name, column_type)
            col_num = 6
            self.__addTable__(doc, talbe_content, col_num)
        doc.save(self.file_path)


        pass


    def __addParagraph__(self, doc, paragraph_content):
        """
        在word文档中添加段落
        :param doc: 文档操作对象
        :param paragraph_content: 段落内容
        :return:
        """
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(paragraph_content)

        # 设置字体格式
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        # 设置字体大小
        run.font.size = Pt(14)
        # 是否加粗
        run.font.bold = True


    def __addTable__(self, doc, table_content, col_num):
        """
        在word文档中添加表格
        :param doc: 文档对象
        :param table_content:表格内容
        :param col_num:列数
        :return:
        """

        #初始化一个空的table
        table = doc.add_table(rows=1, cols=col_num)
        table.style = 'Light Grid'

        """
        表格对齐方式
        WD_TABLE_ALIGNMENT.CENTER 中间对齐
        WD_TABLE_ALIGNMENT.LEFT 左对齐
        WD_TABLE_ALIGNMENT.RIGHT 右对齐
        """
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置列宽
        widths = [1.53,5.11,2.89,1.72]
        self.__set_col_width__(table, widths)

        # 添加表头
        contents = ["序号", "字段名称", "字段代码", "字段类型", "备注", "说明" ]
        self.__addTableHeader__(table,contents)


        table.style = 'Medium Grid 2'

        for tr_list in table_content:
            row_cells = table.add_row().cells
            index = 0
            for td_list in tr_list:
                row_cells[index].text = str(td_list)
                index = index + 1
        pass


    def __set_col_width__(self, table, widths):
        """
        设置列宽
        :param table:表对象
        :param widths:各列的宽度单位默认为cm
        :return:
        """
        for index, width in enumerate(widths):
            table.columns[index].width = Cm(width)
        pass


    def __addTableHeader__(self,table,contents):
        """
        添加表头
        :param table:表对象
        :param contents:表头内容，为一个数组
        :return:
        """

        for index, content in enumerate(contents):
            table.rows[0].cells[index].text = content





if __name__ == '__main__':
    table_info = {'table_name': ['t_1', 't_2'], 'table_comment': ['表1', '表2'], 'column_index': [[1., 2.], [1., 2.]], 'column_name': [['name', 'age'], ['name', 'sex']], 'column_type': [['varchar(255)', 'int(11)'], ['varchar(255)', 'varchar(255)']], 'column_comment': [['', ''], ['', '']]}
    file_path = 'cs.doc'
    pen = MyPen(file_path)
    pen.draw(table_info)
